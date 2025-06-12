import streamlit as st
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI 
from langchain.chains import RetrievalQA
import os
import fitz  # PyMuPDF
import docx
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_xai import ChatXAI
from langchain.chains import ConversationalRetrievalChain
from langchain_core.documents import Document
import base64


xai_api_key = st.secrets["xai_api_key"]
# Upload the document
uploaded_file = st.file_uploader("ðŸ“„ Upload a PDF or DOCX file", type=["pdf", "docx"])

if uploaded_file:
    st.success("âœ… File uploaded successfully!")

    if uploaded_file.name.endswith(".pdf"):
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
    elif uploaded_file.name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        st.warning("Unsupported file type")
        text = ""

    st.text_area("ðŸ“„ Extracted Text", text[:1000])

    # Ask for XAI API Key
    #xai_api_key = st.text_input("xai_api_key")

    if xai_api_key:
        os.environ["xai_api_key"] = xai_api_key  # Optional, if the SDK uses env variable

        # Step 1: Split into chunks
        splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        docs = [Document(page_content=chunk) for chunk in splitter.split_text(text)]

        # Step 2: Embed chunks
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-MiniLM-L6-v2")
        db = FAISS.from_documents(docs, embeddings)

        # Step 3: Connect to XAI LLM
        # If LangChain supports XAI via ChatOpenAI-compatible wrapper:
        llm = ChatXAI(
            temperature=0.3,
            api_key=xai_api_key,
            openai_api_base="https://api.x.ai/v1",   # Replace with actual XAI base URL
            model="grok-3-mini-fast"  # Replace with your actual model name
        )

        # Step 4: Retrieval-based QA
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            retriever=db.as_retriever()
        )

        # Step 5: Ask a question
        user_question = st.text_input("ðŸ’¬ Ask a question about your course:")
        if user_question:
            answer = qa_chain.run(user_question)
            st.markdown("ðŸŽ“ *Answer:* " + answer)





# Load local background image and convert to base64
with open("AI dark.jpg", "rb") as image_file:  # Replace with your file name
    img_bytes = image_file.read()
    img_base64 = base64.b64encode(img_bytes).decode()

# Inject custom CSS using local image
st.markdown("""
    <style>
        .stApp {{
            background-image: url("data:image/png;base64,{img_base64}");
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            color: #ffffff;
        }}

        .css-1v0mbdj, .css-1cpxqw2, .css-1d391kg, .css-ffhzg2 {{
            background-color: rgba(0, 0, 0, 0.5);  /* semi-transparent dark box */
            color: #ffffff !important;
            padding: 20px;
            border-radius: 15px;
        }}

        h1, h2, h3, h4, h5, h6, p, label, .stMarkdown, .stTextInput label {{
            color: #ffffff !important;
        }}

        .stTextInput > div > div > input,
        .stTextArea > div > textarea,
        .stTextInput input {{
            background-color: rgba(255, 255, 255, 0.1);
            color: #ffffff;
            border: 1px solid #ffffff;
        }}

        .stFileUploader {{
            color: #ffffff !important;
        }}
    </style>
""", unsafe_allow_html=True)
